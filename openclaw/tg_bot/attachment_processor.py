"""Shared text-extraction logic for chat attachments.

Used by both the Telegram bot (openclaw/tg_bot/bot.py) and the chat
bridge (sentinel-miniapp-v2/brain_routes.py) to turn an uploaded file
into something the LLM can read. The two callers are different surfaces
of the same shared brain, so they should extract the same way — keeping
this in one module avoids drift.

Supported formats (text extraction):
  .pdf        — pypdf
  .docx       — python-docx
  .xlsx       — openpyxl (cells joined per row, sheets concatenated)
  .txt .md .json .csv .py .js .ts .tsx .html .css .yml .yaml .log .ini .toml
              — direct utf-8 read
  Images      — return a "[image: ...]" marker; the LLM gets a note that
                an image was attached but can't read it directly (vision
                wiring lands later).
  Other       — same marker fallback.

Each extractor returns a STRING:
  - First line: human-readable header (filename + format + key stats)
  - Body:       extracted text
Truncated to MAX_PER_FILE_CHARS so a 200-page PDF can't blow the context.
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_PER_FILE_CHARS = 40_000  # ~10k tokens; safe per-attachment budget

# Lazy-loaded singletons — keep model load out of import time so the bot
# starts fast even when no audio attachments arrive in the first turn.
_WHISPER_MODEL = None
_WHISPER_MODEL_SIZE = "base"   # tiny/base/small/medium/large — base is the
                                # default sweet spot for CPU on Windows.


def _get_whisper():
    global _WHISPER_MODEL
    if _WHISPER_MODEL is not None:
        return _WHISPER_MODEL
    try:
        # On Windows, ctranslate2 (faster-whisper's backend) often loads a
        # second OpenMP runtime that conflicts with whatever the host
        # process already has loaded — produces "OMP: Error #15". Setting
        # KMP_DUPLICATE_LIB_OK=TRUE BEFORE the import is the upstream-
        # recommended workaround for CPU inference. Side effects in
        # practice: none on int8 inference at this scale.
        import os as _os
        _os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")
        from faster_whisper import WhisperModel
        # CPU + int8 — fastest cold-start. Auto-downloads model on first run
        # (cached under HF home; ~75 MB for tiny, ~145 MB for base).
        _WHISPER_MODEL = WhisperModel(
            _WHISPER_MODEL_SIZE, device="cpu", compute_type="int8",
        )
        return _WHISPER_MODEL
    except Exception as e:
        logger.warning("faster-whisper unavailable: %s", e)
        return None

_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".json", ".csv", ".tsv", ".log",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".rs", ".go", ".java", ".kt",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".rb", ".php", ".swift",
    ".html", ".htm", ".css", ".scss", ".sass", ".less",
    ".yml", ".yaml", ".toml", ".ini", ".cfg", ".env",
    ".sh", ".bash", ".zsh", ".ps1", ".bat", ".cmd",
    ".sql", ".graphql", ".proto", ".xml", ".svg",
}
_IMG_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".heic",
             ".bmp", ".avif", ".tiff", ".tif"}
_AUDIO_EXTS = {".mp3", ".m4a", ".aac", ".flac", ".wav", ".opus", ".ogg"}
_VIDEO_EXTS = {".mp4", ".mov", ".mkv", ".webm", ".m4v", ".avi"}
_ARCHIVE_EXTS = {".zip", ".tar", ".gz", ".7z", ".rar"}


def _truncate(s: str) -> str:
    if len(s) <= MAX_PER_FILE_CHARS:
        return s
    cut = MAX_PER_FILE_CHARS - 80
    return s[:cut] + f"\n\n…[truncated; original was {len(s):,} chars]"


def _extract_pdf(path: Path) -> tuple[str, str]:
    """Returns (header, body)."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        n_pages = len(reader.pages)
        parts = []
        for i, page in enumerate(reader.pages, 1):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                parts.append(f"--- Page {i} ---\n{txt.strip()}")
        body = "\n\n".join(parts) if parts else "(no extractable text — image-only PDF?)"
        return (f"PDF · {n_pages} page{'s' if n_pages != 1 else ''}", _truncate(body))
    except Exception as e:
        logger.warning("pdf extract failed for %s: %s", path.name, e)
        return ("PDF (extraction failed)", f"(could not extract: {e})")


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include tables too (common in business docs)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                paras.append(" | ".join(cells))
        body = "\n".join(paras) if paras else "(empty document)"
        return (f"DOCX · {len(paras)} paragraph{'s' if len(paras) != 1 else ''}",
                _truncate(body))
    except Exception as e:
        logger.warning("docx extract failed for %s: %s", path.name, e)
        return ("DOCX (extraction failed)", f"(could not extract: {e})")


def _extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"=== Sheet: {sheet} ===")
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    parts.append(" | ".join("" if c is None else str(c) for c in row))
        body = "\n".join(parts) if parts else "(empty workbook)"
        return (f"XLSX · {len(wb.sheetnames)} sheet{'s' if len(wb.sheetnames) != 1 else ''}",
                _truncate(body))
    except Exception as e:
        logger.warning("xlsx extract failed for %s: %s", path.name, e)
        return ("XLSX (extraction failed)", f"(could not extract: {e})")


def _extract_text(path: Path) -> tuple[str, str]:
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
        return (f"Text · {len(raw.splitlines())} lines · {len(raw):,} chars",
                _truncate(raw))
    except Exception as e:
        return ("Text (read failed)", f"(could not read: {e})")


_CACHED_LM_KEY: str | None = None
_CACHED_VL_MODEL: str | None = None    # populated on first /v1/models probe


def _lm_studio_key() -> str | None:
    """Resolve the LM Studio API key. Process env wins; otherwise we read
    .env.local once and cache. Returns None if no key — vision is then
    skipped silently."""
    import os
    val = os.environ.get("LLM_API_KEY") or os.environ.get("LM_API_TOKEN")
    if val:
        return val
    global _CACHED_LM_KEY
    if _CACHED_LM_KEY is not None:
        return _CACHED_LM_KEY or None
    try:
        from pathlib import Path as _P
        for candidate in (
            _P(r"C:\Users\azfar\metamcp-local\.env.local"),
            _P.home() / "metamcp-local" / ".env.local",
        ):
            if candidate.exists():
                for line in candidate.read_text(encoding="utf-8-sig").splitlines():
                    if line.startswith("LLM_API_KEY="):
                        _CACHED_LM_KEY = line.split("=", 1)[1].strip().strip('"').strip("'")
                        return _CACHED_LM_KEY or None
    except Exception:
        pass
    _CACHED_LM_KEY = ""  # negative cache so we don't re-read on every call
    return None


def _find_vision_model() -> str | None:
    """Find a vision-capable model on the local OpenAI-compatible server.

    Two detection paths, in order:
      1. Name heuristic — a model id advertising "-vl-"/"vision" (LM Studio's
         multi-model setups name vision models this way).
      2. Capability probe — llama.cpp's /props reports modalities.vision. A
         vision-capable Qwen (3.6 27B + --mmproj) does NOT carry "vl" in its
         id, so the name heuristic alone misses it (this was the bug: vision
         was live via --mmproj but reported "unavailable"). When /props says
         vision is on, the single served model (from /v1/models) IS the
         vision model.

    The API key is optional — llama-server only requires it when launched
    with --api-key, so we send it if we have one but don't gate the probe on
    it. Caches the result; clear _CACHED_VL_MODEL=None if you swap models."""
    global _CACHED_VL_MODEL
    if _CACHED_VL_MODEL:
        return _CACHED_VL_MODEL
    key = _lm_studio_key()
    headers = {"Authorization": f"Bearer {key}"} if key else {}
    try:
        import urllib.request, json as _json
        req = urllib.request.Request(
            "http://127.0.0.1:1234/v1/models", headers=headers,
        )
        with urllib.request.urlopen(req, timeout=4) as r:
            data = _json.loads(r.read())
        models = data.get("data", [])
        # 1. name heuristic (multi-model LM Studio)
        for m in models:
            mid = (m.get("id") or "").lower()
            if "-vl-" in mid or "vision" in mid or mid.startswith("vl-") or "/vl-" in mid:
                _CACHED_VL_MODEL = m["id"]
                return _CACHED_VL_MODEL
        # 2. capability probe — llama.cpp --mmproj advertises vision on /props
        try:
            preq = urllib.request.Request(
                "http://127.0.0.1:1234/props", headers=headers,
            )
            with urllib.request.urlopen(preq, timeout=4) as r:
                props = _json.loads(r.read())
            if (props.get("modalities") or {}).get("vision") and models:
                _CACHED_VL_MODEL = models[0]["id"]
                return _CACHED_VL_MODEL
        except Exception as e:
            logger.debug("vision /props probe failed: %s", e)
    except Exception as e:
        logger.debug("vision-model probe failed: %s", e)
    return None


_VISION_PROMPT = (
    "Transcribe and describe this image faithfully and completely — the user "
    "will ask follow-up questions, so capture everything, do not summarise.\n"
    "- Transcribe ALL visible text VERBATIM (exact wording, numbers, times, "
    "codes, punctuation). Do not paraphrase.\n"
    "- If it is a table, roster, schedule, form, or receipt: render it ROW BY "
    "ROW as structured rows (e.g. `Name | Time | Location`), one line per row. "
    "Do not skip rows.\n"
    "- For text that is small, blurry, cut off, or partially covered by an "
    "overlay/legend: give your best-effort reading and mark uncertain "
    "characters with [?] rather than omitting the row. Then note exactly which "
    "region is obscured and by what (e.g. 'rows 5-8 partly under the yellow "
    "STANDBY legend').\n"
    "- Note layout, colours, and colour-coding (often meaningful in rosters/"
    "charts), plus objects/people if present.\n"
    "Report what is actually there, not an interpretation."
)


def _extract_image_vision(path: Path) -> tuple[str, str] | None:
    """Send the image to a vision-capable LM Studio model. Returns
    (header, description) on success, None if no vision model is loaded
    or the call failed. Uses urllib for stdlib-only — keeps this module
    importable from any surface without extra deps."""
    model = _find_vision_model()
    if not model:
        return None
    key = _lm_studio_key()   # optional — only needed if llama-server has --api-key
    try:
        import base64, urllib.request, json as _json
        with open(path, "rb") as fh:
            data = fh.read()
        b64 = base64.b64encode(data).decode("ascii")
        ext = path.suffix.lower().lstrip(".") or "png"
        mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
        payload = {
            "model": model,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "text", "text": _VISION_PROMPT},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                ],
            }],
            # Dense rosters/receipts can need a long verbatim transcription;
            # 1024 truncated them. Low temp = faithful OCR-style reading.
            "max_tokens": 3072,
            "temperature": 0.1,
        }
        _headers = {"Content-Type": "application/json"}
        if key:
            _headers["Authorization"] = f"Bearer {key}"
        req = urllib.request.Request(
            "http://127.0.0.1:1234/v1/chat/completions",
            data=_json.dumps(payload).encode(),
            headers=_headers,
        )
        # 180s timeout — LM Studio queues requests when multiple models
        # are loaded, so the vision call may sit behind a chat_turn on the
        # main model. 90s was too tight in practice; 180s gives margin
        # while still failing reasonably if LM Studio is truly stuck.
        with urllib.request.urlopen(req, timeout=180) as r:
            resp = _json.loads(r.read())
        msg = (resp.get("choices") or [{}])[0].get("message") or {}
        text = (msg.get("content") or "").strip()
        if not text:
            return None
        usage = resp.get("usage") or {}
        head = (f"Image · vision ({model}) · "
                f"{usage.get('prompt_tokens','?')}/{usage.get('completion_tokens','?')} tokens")
        return (head, _truncate(text))
    except Exception as e:
        logger.warning("vision extract failed for %s: %s", path.name, e)
        return None


def _extract_image_ocr(path: Path) -> tuple[str, str] | None:
    """Try OCR on the image. Returns (header, body) on success, None if
    OCR isn't available (tesseract binary missing) — caller can then
    fall back to the [image] marker. Pure best-effort: a failed OCR is
    reported as "no text recognized" rather than crashing the message."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None
    try:
        # pytesseract.get_tesseract_version() raises TesseractNotFoundError
        # if the binary isn't on PATH. We treat that as "OCR unavailable"
        # and let the caller fall back to the [image] marker.
        try:
            version = str(pytesseract.get_tesseract_version())
        except Exception:
            return None
        with Image.open(path) as img:
            text = pytesseract.image_to_string(img) or ""
        text = text.strip()
        if not text:
            return ("Image · OCR found no text",
                    "(tesseract returned an empty string — likely a non-text image)")
        return (f"Image · OCR (tesseract {version}) · {len(text):,} chars",
                _truncate(text))
    except Exception as e:
        logger.warning("OCR failed for %s: %s", path.name, e)
        return None


def _extract_audio_asr(path: Path) -> tuple[str, str] | None:
    """Transcribe an audio file via faster-whisper. Returns (header, body)
    on success, None if faster-whisper isn't loadable — caller falls
    back to the [audio] marker."""
    model = _get_whisper()
    if model is None:
        return None
    try:
        # vad_filter=True trims silence at start/end which both speeds
        # up transcription and reduces hallucinated tokens on quiet
        # bookends. beam_size=1 keeps it fast on CPU; raise later if
        # accuracy matters more than latency.
        segments, info = model.transcribe(
            str(path), beam_size=1, vad_filter=True,
        )
        parts = []
        for seg in segments:
            parts.append(seg.text.strip())
        text = " ".join(p for p in parts if p).strip()
        if not text:
            return ("Audio · no speech detected",
                    "(faster-whisper returned no segments — silent or non-speech audio)")
        head = (f"Audio · transcribed via faster-whisper '{_WHISPER_MODEL_SIZE}' · "
                f"lang={info.language} ({info.language_probability:.0%}) · "
                f"duration={info.duration:.1f}s · {len(text):,} chars")
        return (head, _truncate(text))
    except Exception as e:
        logger.warning("ASR failed for %s: %s", path.name, e)
        return None


def extract_from_file(path: Path | str, display_name: str | None = None) -> str:
    """Extract text from an attached file. Returns a single string ready
    to splice into the user_msg, with a header line + body.

    Falls back to a "[image]" / "[audio]" / "[binary]" marker line for
    anything the LLM can't directly read. Never raises — extraction
    failures degrade to a descriptive note instead.
    """
    p = Path(path)
    name = display_name or p.name
    ext = p.suffix.lower()
    try:
        size = p.stat().st_size
    except OSError:
        size = 0

    if not p.exists() or not p.is_file():
        return f"[Attachment missing: {name}]"

    if ext == ".pdf":
        header, body = _extract_pdf(p)
    elif ext == ".docx":
        header, body = _extract_docx(p)
    elif ext == ".xlsx":
        header, body = _extract_xlsx(p)
    elif ext in _TEXT_EXTS:
        header, body = _extract_text(p)
    elif ext in _IMG_EXTS:
        # Try the loaded VL model first (richest output — describes scene
        # + reads text + reads chart data). If no VL model is loaded,
        # fall back to OCR. If neither, the [image marker] fallback.
        vision = _extract_image_vision(p)
        if vision is not None:
            header, body = vision
        else:
            ocr = _extract_image_ocr(p)
            if ocr is not None:
                header, body = ocr
            else:
                return (f"[Attached image: {name} · {ext.lstrip('.')} · {size:,} bytes]\n"
                        f"(Neither vision nor OCR available. Load a VL model "
                        f"in LM Studio, or install tesseract from "
                        f"https://github.com/UB-Mannheim/tesseract/wiki .)")
    elif ext in _AUDIO_EXTS:
        asr = _extract_audio_asr(p)
        if asr is not None:
            header, body = asr
        else:
            return (f"[Attached audio: {name} · {ext.lstrip('.')} · {size:,} bytes]\n"
                    f"(faster-whisper unavailable. Audio transcription disabled.)")
    elif ext in _VIDEO_EXTS:
        return (f"[Attached video: {name} · {ext.lstrip('.')} · {size:,} bytes]\n"
                f"(Video understanding not yet wired.)")
    elif ext in _ARCHIVE_EXTS:
        return f"[Attached archive: {name} · {ext.lstrip('.')} · {size:,} bytes — contents not unpacked]"
    else:
        # Best-effort: try as text anyway. Often unknown extensions
        # are still UTF-8 text files (configs, scripts, logs).
        try:
            raw = p.read_text(encoding="utf-8")
            if raw and raw.isprintable() or "\n" in raw:
                header = f"Unknown type ({ext or 'no ext'}) · read as text"
                body = _truncate(raw)
            else:
                raise ValueError("non-text")
        except Exception:
            return f"[Attached binary: {name} · {ext or 'unknown'} · {size:,} bytes — not readable as text]"

    return f"[Attached: {name} · {header}]\n{body}"
