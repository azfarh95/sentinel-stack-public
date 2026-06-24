import io
import os

DOCINTEL_ENDPOINT = os.environ.get("DOCINTEL_ENDPOINT", "")
DOCINTEL_KEY      = os.environ.get("DOCINTEL_KEY", "")
_di_available     = bool(DOCINTEL_ENDPOINT and DOCINTEL_KEY)


# ── Azure Document Intelligence ───────────────────────────────────────────────

async def _parse_with_docintel(content: bytes, filename: str) -> str:
    from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
    from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
    from azure.core.credentials import AzureKeyCredential

    name_lower = filename.lower()
    is_invoice = any(kw in name_lower for kw in ("invoice", "inv_", "inv-", "receipt"))
    model_id = "prebuilt-invoice" if is_invoice else "prebuilt-layout"

    async with DocumentIntelligenceClient(
        DOCINTEL_ENDPOINT, AzureKeyCredential(DOCINTEL_KEY)
    ) as client:
        poller = await client.begin_analyze_document(
            model_id,
            AnalyzeDocumentRequest(bytes_source=content),
        )
        result = await poller.result()

    return _fmt_invoice(result) if is_invoice else _fmt_layout(result)


def _fmt_layout(result) -> str:
    lines = []
    for page in result.pages or []:
        lines.append(f"\n--- Page {page.page_number} ---")
        for line in page.lines or []:
            lines.append(line.content)
    for i, table in enumerate(result.tables or [], 1):
        lines.append(f"\n[Table {i}]")
        rows: dict = {}
        for cell in table.cells:
            rows.setdefault(cell.row_index, {})[cell.column_index] = cell.content or ""
        for r in sorted(rows):
            col_count = max(rows[r].keys()) + 1
            lines.append("\t".join(rows[r].get(c, "") for c in range(col_count)))
    return "\n".join(lines).strip()


def _fmt_invoice(result) -> str:
    lines = ["=== INVOICE (Azure Document Intelligence) ===\n"]
    for doc in result.documents or []:
        fields = doc.fields or {}
        for name in (
            "VendorName", "CustomerName", "InvoiceId", "InvoiceDate",
            "DueDate", "PurchaseOrder", "SubTotal", "TotalTax",
            "InvoiceTotal", "AmountDue", "BillingAddress", "ShippingAddress",
        ):
            f = fields.get(name)
            if f and f.content:
                lines.append(f"{name}: {f.content}")
        items_field = fields.get("Items")
        if items_field and items_field.value_array:
            lines.append("\nLine Items:")
            for item in items_field.value_array:
                obj = item.value_object or {}
                parts = []
                for key in ("Description", "ProductCode"):
                    v = obj.get(key)
                    if v and v.content:
                        parts.append(v.content)
                        break
                for key, label in (("Quantity", "qty:"), ("UnitPrice", "@"), ("Amount", "= ")):
                    v = obj.get(key)
                    if v and v.content:
                        parts.append(f"{label}{v.content}")
                lines.append("  • " + "  ".join(parts))
    return "\n".join(lines)


# ── Local fallback parsers ────────────────────────────────────────────────────

def _parse_excel(content: bytes, filename: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append("\t".join("" if c is None else str(c) for c in row))
        if rows:
            parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts) if parts else "(empty workbook)"


def _parse_pdf_local(content: bytes) -> str:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    pages = [doc[i].get_text() for i in range(len(doc))]
    doc.close()
    text = "\n\n".join(pages).strip()
    return text if text else "(no extractable text — may be a scanned image PDF)"


def _parse_word_local(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines) if lines else "(empty document)"


# ── Public entry point (async) ────────────────────────────────────────────────

async def parse_file(content: bytes, filename: str, mime_type: str = "") -> str:
    name = filename.lower()
    try:
        if name.endswith((".xlsx", ".xls", ".xlsm", ".xlsb")):
            return _parse_excel(content, filename)

        if name.endswith((".csv", ".txt", ".md", ".json", ".xml", ".html", ".htm")):
            return content.decode("utf-8", errors="replace")

        if name.endswith(".pdf") or "pdf" in mime_type or name.endswith((".docx", ".doc")):
            if _di_available:
                return await _parse_with_docintel(content, filename)
            if name.endswith(".pdf") or "pdf" in mime_type:
                return _parse_pdf_local(content)
            return _parse_word_local(content)

        return (
            f"(unsupported format — {len(content):,} bytes. "
            "Supported: xlsx, pdf, docx, csv, txt, md, json)"
        )
    except Exception as e:
        return f"(parse error: {e})"
